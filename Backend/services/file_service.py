"""
文件处理服务 - 负责文件上传、内容提取和分块处理
"""
import os
import io
import uuid
import hashlib
import logging
from datetime import datetime, timezone
from typing import List, Dict, Optional, Tuple
from werkzeug.datastructures import FileStorage
import chardet
from pypdf import PdfReader
from docx import Document
from PIL import Image

from ..models import db, KnowledgeItem
from ..services.knowledge_service import KnowledgeService
from ..services.vector_service import VectorService
from ..utils.text_utils import clean_text, extract_keywords

logger = logging.getLogger(__name__)


class FileService:
    """文件处理服务类"""
    
    # 支持的文件类型
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.pdf': 'application/pdf',
        '.doc': 'application/msword',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.rtf': 'application/rtf'
    }
    
    # 最大文件大小 (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    def __init__(self):
        self._knowledge_service = None
        self._vector_service = None
        self.upload_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'uploads')
        self._ensure_upload_dir()
    
    @property
    def knowledge_service(self):
        """懒加载知识库服务"""
        if self._knowledge_service is None:
            self._knowledge_service = KnowledgeService()
        return self._knowledge_service
    
    @property
    def vector_service(self):
        """懒加载向量服务（避免不必要的模型加载）"""
        if self._vector_service is None:
            self._vector_service = VectorService()
        return self._vector_service
    
    def _ensure_upload_dir(self):
        """确保上传目录存在"""
        if not os.path.exists(self.upload_dir):
            os.makedirs(self.upload_dir)
    
    def validate_file(self, file: FileStorage) -> Dict:
        """
        验证上传的文件
        
        Args:
            file: 上传的文件对象
            
        Returns:
            Dict: 验证结果
        """
        if not file:
            return {
                'valid': False,
                'error': '没有上传文件'
            }
        
        # 检查文件名
        if not file.filename:
            return {
                'valid': False,
                'error': '文件名不能为空'
            }
        
        # 检查文件扩展名
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return {
                'valid': False,
                'error': f'不支持的文件类型: {file_ext}。支持的类型: {", ".join(self.SUPPORTED_EXTENSIONS.keys())}'
            }
        
        # 检查文件大小
        file.seek(0, 2)  # 移动到文件末尾
        file_size = file.tell()
        file.seek(0)  # 重置到文件开头
        
        if file_size > self.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f'文件大小超过限制: {file_size} bytes > {self.MAX_FILE_SIZE} bytes'
            }
        
        return {
            'valid': True,
            'filename': file.filename,
            'extension': file_ext,
            'size': file_size,
            'mime_type': self.SUPPORTED_EXTENSIONS[file_ext]
        }
    
    def upload_file(self, file: FileStorage, category: str = None, tags: List[str] = None) -> Dict:
        """
        上传并处理文件
        
        Args:
            file: 上传的文件对象
            category: 文件分类
            tags: 文件标签
            
        Returns:
            Dict: 处理结果
        """
        try:
            # 验证文件
            validation = self.validate_file(file)
            if not validation['valid']:
                return {
                    'success': False,
                    'error': validation['error']
                }
            
            # 生成唯一文件名
            file_id = str(uuid.uuid4())
            file_ext = validation['extension']
            saved_filename = f"{file_id}{file_ext}"
            file_path = os.path.join(self.upload_dir, saved_filename)
            
            # 保存文件
            file.save(file_path)
            
            logger.info(f"文件上传成功: {file.filename} -> {saved_filename}")
            
            # 提取文件内容
            extraction_result = self.extract_file_content(file_path, file_ext)
            
            if not extraction_result['success']:
                # 删除已保存的文件
                os.remove(file_path)
                return extraction_result
            
            content = extraction_result['content']
            if not content or len(content.strip()) < 10:
                os.remove(file_path)
                return {
                    'success': False,
                    'error': '文件内容为空或太短，无法处理'
                }
            
            # 分块处理内容
            chunks = self.chunk_text(content, chunk_size=1000, overlap=200)
            
            if not chunks:
                os.remove(file_path)
                return {
                    'success': False,
                    'error': '无法将内容分块处理'
                }
            
            # 保存到知识库
            saved_count = 0
            knowledge_items = []
            
            for i, chunk in enumerate(chunks):
                try:
                    # 生成内容哈希
                    content_hash = self._generate_content_hash(chunk)
                    
                    # 检查是否已存在
                    existing_item = KnowledgeItem.query.filter_by(content_hash=content_hash).first()
                    if existing_item:
                        continue
                    
                    # 创建知识库条目
                    title = f"{os.path.splitext(file.filename)[0]} - 第{i+1}部分"
                    
                    knowledge_item = KnowledgeItem(
                        title=title,
                        content=chunk,
                        source_url=f"file://{saved_filename}",
                        source_name=file.filename,
                        source_type='upload',
                        category=category or 'document',
                        tags=tags or [],
                        content_hash=content_hash,
                        created_at=datetime.now(timezone.utc)
                    )
                    
                    db.session.add(knowledge_item)
                    db.session.flush()  # 获取ID
                    
                    # 向量化
                    try:
                        vector_id = self.vector_service.add_document(
                            knowledge_item.id,
                            title + ' ' + chunk
                        )
                        knowledge_item.vector_id = vector_id
                        saved_count += 1
                        knowledge_items.append(knowledge_item.to_dict())
                        
                    except Exception as e:
                        logger.error(f"向量化失败: {str(e)}")
                        # 即使向量化失败，也保存文本内容
                        saved_count += 1
                        knowledge_items.append(knowledge_item.to_dict())
                
                except Exception as e:
                    logger.error(f"保存知识库条目失败: {str(e)}")
                    continue
            
            db.session.commit()
            
            return {
                'success': True,
                'message': f'文件处理完成，保存了 {saved_count} 个知识库条目',
                'filename': file.filename,
                'file_id': file_id,
                'chunks_count': len(chunks),
                'saved_count': saved_count,
                'knowledge_items': knowledge_items
            }
            
        except Exception as e:
            logger.error(f"文件处理失败: {str(e)}")
            return {
                'success': False,
                'error': f'文件处理失败: {str(e)}'
            }
    
    def extract_file_content(self, file_path: str, file_ext: str) -> Dict:
        """
        提取文件内容
        
        Args:
            file_path: 文件路径
            file_ext: 文件扩展名
            
        Returns:
            Dict: 提取结果
        """
        try:
            content = ""
            
            if file_ext == '.txt':
                content = self._extract_txt_content(file_path)
            elif file_ext == '.md':
                content = self._extract_markdown_content(file_path)
            elif file_ext == '.pdf':
                content = self._extract_pdf_content(file_path)
            elif file_ext in ['.doc', '.docx']:
                content = self._extract_doc_content(file_path)
            elif file_ext == '.rtf':
                content = self._extract_rtf_content(file_path)
            else:
                return {
                    'success': False,
                    'error': f'不支持的文件类型: {file_ext}'
                }
            
            if content:
                content = clean_text(content)
                return {
                    'success': True,
                    'content': content,
                    'content_length': len(content)
                }
            else:
                return {
                    'success': False,
                    'error': '无法提取文件内容'
                }
                
        except Exception as e:
            logger.error(f"提取文件内容失败: {str(e)}")
            return {
                'success': False,
                'error': f'提取文件内容失败: {str(e)}'
            }
    
    def _extract_txt_content(self, file_path: str) -> str:
        """提取TXT文件内容"""
        try:
            # 检测文件编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # 尝试不同的编码
            encodings = [encoding, 'utf-8', 'gbk', 'gb2312', 'big5']
            
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误处理
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"提取TXT内容失败: {str(e)}")
            return ""
    
    def _extract_markdown_content(self, file_path: str) -> str:
        """提取Markdown文件内容"""
        try:
            # Markdown 文件与 TXT 文件处理方式相同，都是纯文本
            # 检测文件编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # 尝试不同的编码
            encodings = [encoding, 'utf-8', 'gbk', 'gb2312', 'big5']
            
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误处理
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"提取Markdown内容失败: {str(e)}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """提取PDF文件内容"""
        try:
            content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
            
            return content
            
        except Exception as e:
            logger.error(f"提取PDF内容失败: {str(e)}")
            return ""
    
    def _extract_doc_content(self, file_path: str) -> str:
        """提取DOC/DOCX文件内容"""
        try:
            content = ""
            
            # 读取Word文档
            doc = Document(file_path)
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += cell.text + "\t"
                    content += "\n"
            
            return content
            
        except Exception as e:
            logger.error(f"提取DOC/DOCX内容失败: {str(e)}")
            return ""
    
    def _extract_rtf_content(self, file_path: str) -> str:
        """提取RTF文件内容"""
        try:
            # RTF文件处理比较复杂，这里简化处理
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # 简单的RTF标签移除
            import re
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)  # 移除RTF控制字符
            content = re.sub(r'[{}]', '', content)  # 移除大括号
            content = re.sub(r'\s+', ' ', content)  # 合并空白字符
            
            return content
            
        except Exception as e:
            logger.error(f"提取RTF内容失败: {str(e)}")
            return ""
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        将文本分块处理 - 使用滑动窗口算法，O(n)时间复杂度
        
        算法特点：
        1. 固定步长滑动，避免死循环
        2. 智能边界分割，优先在句子边界处分割
        3. 内存高效，只创建必要的字符串切片
        4. 防御性编程，处理各种边界情况
        
        Args:
            text: 要分块的文本
            chunk_size: 每块的大小（字符数）
            overlap: 重叠的字符数
            
        Returns:
            List[str]: 文本块列表
        """
        # 边界情况处理
        if not text:
            return []
        
        text = text.strip()
        if not text:
            return []
        
        text_len = len(text)
        if text_len <= chunk_size:
            return [text]
        
        # 参数验证和自动调整
        chunk_size = max(100, chunk_size)  # 最小块大小100字符
        overlap = max(0, min(overlap, chunk_size // 2))  # overlap不超过chunk_size的一半
        
        # 计算固定步长（关键：避免死循环）
        step_size = chunk_size - overlap
        if step_size <= 0:
            step_size = chunk_size // 2  # 保底策略
        
        chunks = []
        start = 0
        
        # 句子分隔符（按优先级排序）
        sentence_delimiters = ['。', '！', '？', '\n', '；', '，', '.', '!', '?', ';']
        
        while start < text_len:
            # 计算当前块的结束位置
            end = min(start + chunk_size, text_len)
            
            # 智能边界调整：在句子边界处分割
            if end < text_len:  # 不是最后一块
                # 在最后20%的范围内查找分隔符
                search_start = max(start, end - chunk_size // 5)
                best_split = -1
                
                # 按优先级查找分隔符
                for delimiter in sentence_delimiters:
                    pos = text.rfind(delimiter, search_start, end)
                    if pos != -1:
                        best_split = pos + 1  # 包含分隔符
                        break
                
                # 如果找到了合适的分割点
                if best_split > start:
                    end = best_split
            
            # 提取块（避免重复切片）
            chunk = text[start:end].strip()
            
            # 只添加非空块
            if chunk:
                chunks.append(chunk)
            
            # 固定步长前进（关键：保证O(n)复杂度）
            start += step_size
            
            # 防御性检查：如果start没有前进，强制前进
            if start <= (end - step_size):
                start = end
        
        return chunks
    
    def _generate_content_hash(self, content: str) -> str:
        """生成内容哈希值"""
        return hashlib.md5(content.encode('utf-8')).hexdigest()
    
    def get_uploaded_files_info(self) -> Dict:
        """获取已上传文件的信息"""
        try:
            # 从知识库中获取上传的文件信息
            uploaded_items = KnowledgeItem.query.filter_by(source_type='upload').all()
            
            files_info = {}
            
            for item in uploaded_items:
                source_name = item.source_name
                if source_name not in files_info:
                    files_info[source_name] = {
                        'filename': source_name,
                        'category': item.category,
                        'tags': item.tags,
                        'created_at': item.created_at.isoformat(),
                        'chunks_count': 0,
                        'total_content_length': 0
                    }
                
                files_info[source_name]['chunks_count'] += 1
                files_info[source_name]['total_content_length'] += len(item.content)
            
            return {
                'success': True,
                'files_count': len(files_info),
                'files': list(files_info.values())
            }
            
        except Exception as e:
            logger.error(f"获取上传文件信息失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def delete_uploaded_file(self, filename: str) -> Dict:
        """
        删除上传的文件及其相关数据
        
        Args:
            filename: 文件名
            
        Returns:
            Dict: 删除结果
        """
        try:
            # 删除知识库中的相关条目
            deleted_items = KnowledgeItem.query.filter_by(
                source_name=filename,
                source_type='upload'
            ).all()
            
            deleted_count = len(deleted_items)
            
            for item in deleted_items:
                # 从向量数据库中删除
                if item.vector_id:
                    try:
                        self.vector_service.delete_document(item.vector_id)
                    except Exception as e:
                        logger.error(f"删除向量数据失败: {str(e)}")
                
                db.session.delete(item)
            
            db.session.commit()
            
            # 删除物理文件
            file_pattern = filename.split('.')[0]  # 获取文件名（不含扩展名）
            upload_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'uploads')
            
            for file_name in os.listdir(upload_dir):
                if file_name.startswith(file_pattern):
                    file_path = os.path.join(upload_dir, file_name)
                    try:
                        os.remove(file_path)
                        logger.info(f"删除物理文件: {file_path}")
                    except Exception as e:
                        logger.error(f"删除物理文件失败: {str(e)}")
            
            return {
                'success': True,
                'message': f'成功删除文件 {filename} 及其 {deleted_count} 个知识库条目',
                'deleted_items_count': deleted_count
            }
            
        except Exception as e:
            logger.error(f"删除上传文件失败: {str(e)}")
            db.session.rollback()
            return {
                'success': False,
                'error': str(e)
            }
